import docx
import requests
import re
import os
import matplotlib.pyplot as plot
from PIL import Image
from docx.shared import Cm,Inches,RGBColor

""" project link: https://www.gutenberg.org/cache/epub/67238/pg67238.txt
    Title: The Cricket on the Hearth A Fairy Tale of Home
    Author: Charles Dickens """


def run():
    print("="*70 + " Welcome Lab 11 "+"="*70)
    book = requests.get('https://www.gutenberg.org/cache/epub/67238/pg67238.txt')
    book_content = book.text.strip()

    """ Task 2 """
    title = extract_title(book_content)
    print('Title of the book is "'+title+'"')
    author = extract_author(book_content)
    print('Author name is '+author)
    #print(extract_chapter(book_content))

    """ Task 3 """
    words_count = num_of_words(book_content)
    print("\nTotal words in chapter I is", words_count)
    chapterOne = extract_chapter(book_content, 'CHAPTER I', 'CHAPTER II')
    paragraphs = paragraph_in_chapter(chapterOne)
    print(paragraphs)
    word_counted = words_in_paragraph(paragraphs)
    print(word_counted)

    plot.figure(figsize=(7,5))
    plot.title("Length of Paragraph")
    plot.pie(word_counted)
    plot.savefig('distribution.png')
    plot.show()

    """ Task 4 """
    download_cover_photo()
    """ Task 5 """
    crop_and_resize()
    """ Task 6 """
    black_and_white_image()

    """ Task 7 """
    generate_document(book_content, chapterOne, paragraphs, word_counted)

def download_cover_photo():
    image = requests.get('https://encrypted-tbn3.gstatic.com/images?q=tbn:ANd9GcS-AERAYC3PRxvW2cXwKgjK5kjwJ6xhwtOhU8vHDtQmzb3n6sUs')
    with open('CoverImage.jpg','wb') as file:
        file.write(image.content)


def round_ten(value):
    return value - (value%10)

def words_in_paragraph(paragraphs):
    words_array = []
    for words in paragraphs:
        words_array.append(round_ten(len(words.replace('\n', " ").split(" "))))
    return sorted(words_array)

def crop_and_resize():
    img = Image.open('CoverImage.jpg')
    img2 = img.crop((100,100,400,200))
    img2.save('crop_image.jpg')
    w, h = img.size
    img3 = img2.resize((w+100, h+50))
    img3.save('resize_image.jpg')

def paragraph_in_chapter(chapter_one):
    paragraph_array = []
    paragraphs = chapter_one.replace('\r', '').split('\n\n')
    for text in paragraphs:
        if text != "" and text != '\n':
            paragraph_array.append(text)
    return paragraph_array


def num_of_words(content):
    length = len(content.split())
    return length - (length%10)


def extract_title(content):
    title_line = re.search('Title: ', content)
    print(title_line)
    return content[title_line.span()[1]: content.find('Author', title_line.span()[0])].strip()


def generate_document(content, chapterOne, paragraphs, words):
    report = docx.Document()
    report.add_heading(extract_title(content), 0)
    report.add_heading('Author: '+ extract_author(content), 1)
    report.add_heading('Author of the Report: Md Shahadat Hossen Nayem', 2)
    report.add_picture('CoverImage.jpg', width=Inches(3))
    report.add_page_break()

    report.add_heading('Information Page', 1)
    info = report.add_paragraph()
    report.add_picture('distribution.png', width=Inches(4))
    run = info.add_run('Length Distribution Plot\n')
    font = run.font
    font.color.rgb = RGBColor(255, 0, 0)
    info.add_run('The plot is describing Chapter I\n').bold = True
    info.add_run('Total Length of Chapter I: '+ str(len(chapterOne)) + '\n').italic = True
    info.add_run('Number of paragraph in Chapter I:'+ str(len(paragraphs)) + '\n').italic = True
    info.add_run('Number of words in Chapter I:'+ str(len(words)) + '\n').italic = True
    info.add_run('Words in Longest paragraph in Chapter I:'+ str(max(words))+ '\n').underline = True
    info.add_run('Words in Shortest paragraph in Chapter I:'+ str(min(words)) + '\n').underline = True

    report.save('Report.docx')
    os.system('start Report.docx')


def extract_author(content):
    title_line = re.search('Author:', content)
    return content[title_line.span()[1]: content.find('\n',title_line.span()[0])].strip()

def black_and_white_image():
    img = requests.get('https://i.pinimg.com/236x/5c/30/63/5c3063e5f763e51bc667e1f2fbb0112f--panda-logo-design-love.jpg')
    with open('Black&White','wb') as file:
        file.write(img.content)
    logo = Image.open('Black&White')
    main_image = Image.open('CoverImage.jpg')
    location = (20, 30)
    main_image.paste(logo, location)
    main_image.save('Cover&Logo.jpg')


def extract_chapter(bookText, beginning, end):
    try:
        startInd = bookText.index(beginning) + len(beginning)
        endInd = bookText.index(end, startInd)
        return bookText[startInd:endInd]
    except ValueError:
        return ""


if __name__ == '__main__':
    run()
